from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ── Helper: add image placeholder ──
def add_image_placeholder(doc, caption, width=Inches(5), height=Inches(3)):
    """Add a bordered empty rectangle with a caption as screenshot placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)
    # Create a table cell with a border to simulate an image frame
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = width
    cell.height = height
    # Set cell shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    p_cell = cell.paragraphs[0]
    p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cell = p_cell.add_run(f"[Capture d'écran : {caption}]")
    run_cell.font.size = Pt(11)
    run_cell.font.color.rgb = RGBColor(120, 120, 120)
    run_cell.font.italic = True
    # Remove table borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)
    doc.add_paragraph()  # spacing
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(80, 80, 80)

def heading(level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def paragraph(text, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RAPPORT DE STAGE")
r.font.size = Pt(26)
r.font.bold = True
r.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Réalisation d'une plateforme de gestion\npour l'espace professeurs")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0, 51, 102)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Présenté par : Hamza Linawwi")
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Encadré par : Mehdi Sekale")
r.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Année universitaire : 2025/2026")
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
heading(1, "Table des matières")
toc_items = [
    ("Introduction générale", 3),
    ("Chapitre 1 : Contexte et présentation de l'organisme d'accueil", 4),
    ("  1.1 Présentation de l'établissement", 4),
    ("  1.2 Organigramme", 4),
    ("  1.3 Description du service", 5),
    ("Chapitre 2 : Analyse et conception du projet", 6),
    ("  2.1 Problématique et objectifs", 6),
    ("  2.2 Analyse des besoins", 6),
    ("  2.3 Diagrammes UML", 7),
    ("  2.4 Architecture technique", 8),
    ("Chapitre 3 : Technologies et outils utilisés", 9),
    ("  3.1 Environnement de développement", 9),
    ("  3.2 Technologies front-end", 9),
    ("  3.3 Technologies back-end", 10),
    ("  3.4 Base de données", 10),
    ("Chapitre 4 : Réalisation et implémentation", 11),
    ("  4.1 Modules développés", 11),
    ("  4.2 Interfaces utilisateur", 12),
    ("  4.3 Schéma de la base de données", 14),
    ("Conclusion générale", 15),
    ("Bibliographie et webographie", 16),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"{item}")
    r.font.size = Pt(12)
    tab = p.add_run("\t")
    r2 = p.add_run(f"{page}")
    r2.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════
heading(1, "Introduction générale")
paragraph(
    "Dans le cadre de ma formation, j'ai effectué un stage au sein de l'établissement "
    "Chariaa, ayant pour objectif la conception et le développement d'une plateforme "
    "web dédiée à la gestion de l'espace professeurs. Ce projet vise à moderniser et "
    "simplifier la gestion administrative des enseignants, en leur offrant des outils "
    "numériques performants pour le suivi des étudiants, la gestion des notes, "
    "l'attribution des modules et l'organisation des examens."
)
paragraph(
    "Ce rapport présente l'ensemble des travaux réalisés durant cette période de stage. "
    "Il est structuré en quatre chapitres : le premier présente l'organisme d'accueil, "
    "le deuxième détaille l'analyse et la conception du projet, le troisième décrit "
    "les technologies utilisées, et le dernier chapitre expose la réalisation et "
    "l'implémentation de la plateforme."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 1
# ══════════════════════════════════════════════════════════════
heading(1, "Chapitre 1 : Contexte et présentation de l'organisme d'accueil")

heading(2, "1.1 Présentation de l'établissement")
paragraph(
    "L'établissement Chariaa est un établissement d'enseignement supérieur qui offre "
    "une gamme variée de formations dans différents domaines. Il accueille chaque année "
    "un grand nombre d'étudiants et dispose d'un corps professoral qualifié. "
    "L'établissement s'engage à offrir une éducation de qualité et à former des "
    "professionnels compétents capables de répondre aux besoins du marché du travail."
)
paragraph(
    "Doté d'une infrastructure moderne, l'établissement comprend plusieurs départements, "
    "des salles de cours équipées, des laboratoires informatiques, une bibliothèque, "
    "ainsi que des espaces administratifs dédiés à la gestion pédagogique."
)

add_image_placeholder(doc, "Logo de l'établissement Chariaa", width=Inches(3), height=Inches(2))
add_image_placeholder(doc, "Vue d'ensemble de l'établissement", width=Inches(5), height=Inches(2.5))

heading(2, "1.2 Organigramme")
paragraph(
    "L'organigramme ci-dessous présente la structure hiérarchique de l'établissement."
)
add_image_placeholder(doc, "Organigramme de l'établissement", width=Inches(5), height=Inches(3))

heading(2, "1.3 Description du service")
paragraph(
    "Le stage s'est déroulé au sein du service informatique et pédagogique de "
    "l'établissement. Ce service est responsable de la gestion des systèmes "
    "d'information, du suivi des inscriptions pédagogiques, de la gestion des "
    "emplois du temps, et de l'assistance technique aux utilisateurs. "
    "L'équipe est composée d'ingénieurs et de techniciens spécialisés."
)
paragraph(
    "C'est dans ce contexte que le besoin d'une plateforme dédiée à l'espace "
    "professeurs a émergé, afin de dématérialiser et centraliser les différentes "
    "tâches administratives liées à l'enseignement."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 2
# ══════════════════════════════════════════════════════════════
heading(1, "Chapitre 2 : Analyse et conception du projet")

heading(2, "2.1 Problématique et objectifs")
paragraph(
    "Actuellement, la gestion des activités pédagogiques des professeurs est "
    "effectuée de manière manuelle ou via des outils disparates, ce qui engendre :"
)
bullet("Une perte de temps dans les tâches administratives")
bullet("Des risques d'erreurs dans la saisie des notes")
bullet("Une difficulté à assurer le suivi des inscriptions pédagogiques")
bullet("Un manque de centralisation des informations")
paragraph(
    "L'objectif principal de ce projet est de concevoir et développer une plateforme "
    "web intitulée « Prof Space Chariaa » qui permettra de :"
)
bullet("Gérer les inscriptions pédagogiques des étudiants")
bullet("Permettre aux professeurs de saisir et consulter les notes d'examens")
bullet("Assurer la répartition des étudiants dans les salles d'examen")
bullet("Faciliter la gestion des modules et des groupes")
bullet("Offrir une interface d'administration pour la gestion des utilisateurs")
bullet("Exporter les données vers Excel pour un traitement ultérieur")

heading(2, "2.2 Analyse des besoins")
heading(3, "2.2.1 Besoins fonctionnels")
paragraph("Les besoins fonctionnels identifiés sont les suivants :")
bullet("Authentification et gestion des rôles (Admin, Super Admin, Professeur)")
bullet("Gestion des étudiants (CRUD, import/export Excel, photo)")
bullet("Gestion des modules et des groupes")
bullet("Gestion des inscriptions pédagogiques")
bullet("Saisie des notes d'examens avec calcul automatique des décisions")
bullet("Répartition des étudiants dans les salles d'examen")
bullet("Gestion des professeurs et attribution des modules")
bullet("Interface multilingue (français / arabe)")

heading(3, "2.2.2 Besoins non fonctionnels")
bullet("Interface responsive et intuitive")
bullet("Sécurisation des accès via authentification et rôles")
bullet("Performance et rapidité d'exécution")
bullet("Maintenabilité et évolutivité du code")
bullet("Base de données relationnelle normalisée")

heading(2, "2.3 Diagrammes UML")

heading(3, "2.3.1 Diagramme de cas d'utilisation")
paragraph(
    "Le diagramme de cas d'utilisation ci-dessous illustre les interactions entre "
    "les différents acteurs (Administrateur, Professeur) et le système."
)
add_image_placeholder(doc, "Diagramme de cas d'utilisation", width=Inches(5), height=Inches(3.5))

heading(3, "2.3.2 Diagramme de classes")
paragraph(
    "Le diagramme de classes présente la structure statique du système, avec les "
    "entités principales et leurs relations."
)
add_image_placeholder(doc, "Diagramme de classes", width=Inches(5.5), height=Inches(3.5))

heading(3, "2.3.3 Diagramme de séquence")
paragraph(
    "Le diagramme de séquence ci-dessous illustre le flux d'interaction pour le "
    "processus de saisie des notes par un professeur."
)
add_image_placeholder(doc, "Diagramme de séquence - Saisie des notes", width=Inches(5.5), height=Inches(3))

heading(2, "2.4 Architecture technique")
paragraph(
    "L'application est construite selon une architecture moderne :"
)
bullet("Architecture MVC (Modèle-Vue-Contrôleur) côté serveur avec Laravel")
bullet("Front-end en React avec Inertia.js pour une expérience SPA sans perdre les avantages du SSR")
bullet("API RESTful pour les opérations asynchrones (import/export, téléchargement de photos)")
bullet("Base de données MySQL avec schéma normalisé")
bullet("Authentification via Laravel Breeze avec gestion des rôles")

add_image_placeholder(doc, "Architecture technique de l'application", width=Inches(5), height=Inches(3))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 3
# ══════════════════════════════════════════════════════════════
heading(1, "Chapitre 3 : Technologies et outils utilisés")

heading(2, "3.1 Environnement de développement")
paragraph("Le développement de l'application a été réalisé dans l'environnement suivant :")
bullet("Système d'exploitation : Windows 10/11")
bullet("IDE : Visual Studio Code")
bullet("Serveur local : Laragon / XAMPP")
bullet("Gestionnaire de versions : Git")
bullet("Hébergement de code : GitHub")

heading(2, "3.2 Technologies front-end")
paragraph("Pour la partie front-end, les technologies suivantes ont été utilisées :")

p = doc.add_paragraph()
r = p.add_run("React.js")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Bibliothèque JavaScript pour la construction d'interfaces utilisateur réactives et dynamiques.")

p = doc.add_paragraph()
r = p.add_run("Inertia.js")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Permet de construire des applications monopages (SPA) en utilisant des composants React sans avoir besoin d'API REST.")

p = doc.add_paragraph()
r = p.add_run("Tailwind CSS")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Framework CSS utilitaire pour un design moderne et responsive.")

paragraph(
    "L'interface est conçue pour être entièrement responsive et supporte le mode "
    "RTL (Right-to-Left) pour la localisation arabe."
)

heading(2, "3.3 Technologies back-end")

p = doc.add_paragraph()
r = p.add_run("Laravel")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Framework PHP MVC puissant offrant une vaste gamme de fonctionnalités (ORM Eloquent, migrations, validation, etc.).")

p = doc.add_paragraph()
r = p.add_run("PHP 8.2")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Langage de programmation côté serveur.")

p = doc.add_paragraph()
r = p.add_run("MySQL")
r.font.bold = True
r.font.size = Pt(12)
p.add_run(" : Système de gestion de bases de données relationnelles.")

heading(2, "3.4 Base de données")
paragraph(
    "La base de données est structurée autour des entités principales suivantes : "
    "utilisateurs (users), professeurs (prof), étudiants (etudiant), modules (module), "
    "inscriptions pédagogiques (etudiant_module), notes d'examens (note_exam), "
    "salles (salle), groupes (groupes), niveaux (niveaux), semestres (semestres), "
    "et filières (filieres). Le schéma relationnel assure l'intégrité des données "
    "via des clés étrangères et des contraintes d'unicité."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 4
# ══════════════════════════════════════════════════════════════
heading(1, "Chapitre 4 : Réalisation et implémentation")

heading(2, "4.1 Modules développés")

heading(3, "4.1.1 Module d'authentification et de gestion des rôles")
paragraph(
    "Le système d'authentification utilise Laravel Breeze. Trois rôles sont définis : "
    "Admin, Super Admin et Professeur. Chaque rôle a des permissions spécifiques "
    "qui déterminent les fonctionnalités accessibles."
)
add_image_placeholder(doc, "Page de connexion", width=Inches(5), height=Inches(2.5))

heading(3, "4.1.2 Module de gestion des étudiants")
paragraph(
    "Ce module permet d'effectuer toutes les opérations CRUD sur les étudiants, "
    "avec des fonctionnalités avancées comme :"
)
bullet("Importation massive depuis un fichier Excel avec validation")
bullet("Exportation filtrée vers Excel ou CSV")
bullet("Téléchargement et suppression de photo de profil")
bullet("Recherche et filtres (sexe, filière)")
bullet("Affichage en mode grille ou liste")
add_image_placeholder(doc, "Liste des étudiants - mode grille", width=Inches(5), height=Inches(3))
add_image_placeholder(doc, "Fiche détaillée d'un étudiant", width=Inches(5), height=Inches(3))
add_image_placeholder(doc, "Fenêtre d'import Excel", width=Inches(5), height=Inches(2.5))
add_image_placeholder(doc, "Fenêtre d'export avec sélection des champs", width=Inches(5), height=Inches(2.5))

heading(3, "4.1.3 Module de gestion des professeurs")
paragraph(
    "Ce module permet la gestion des comptes professeurs et l'attribution des "
    "modules d'enseignement. Chaque professeur peut se voir attribuer un ou "
    "plusieurs modules. L'interface d'attribution affiche les modules disponibles "
    "avec leur hiérarchie (filière → niveau → semestre → module) et le coefficient."
)
add_image_placeholder(doc, "Liste des professeurs", width=Inches(5), height=Inches(2.5))
add_image_placeholder(doc, "Attribution des modules à un professeur", width=Inches(5), height=Inches(3))
add_image_placeholder(doc, "Modules enseignés par un professeur", width=Inches(5), height=Inches(2.5))

heading(3, "4.1.4 Module de gestion des inscriptions pédagogiques")
paragraph(
    "Ce module gère l'inscription des étudiants aux modules. Il permet d'ajouter "
    "ou de retirer des inscriptions et d'afficher les statistiques par module, "
    "niveau, semestre et filière."
)
add_image_placeholder(doc, "Gestion des inscriptions pédagogiques", width=Inches(5), height=Inches(3))

heading(3, "4.1.5 Module de saisie des notes d'examens")
paragraph(
    "Les professeurs peuvent saisir les notes normales et de rattrapage pour "
    "chaque étudiant inscrit à leurs modules. Le système calcule automatiquement :"
)
bullet("La note finale (la meilleure entre note normale et note de rattrapage)")
bullet("Les décisions en français et en arabe (Validé / Non validé)")
bullet("Le statut de l'examen (normale, rattrapage, finale)")
add_image_placeholder(doc, "Saisie des notes d'examens", width=Inches(5), height=Inches(3))
add_image_placeholder(doc, "Calcul automatique des décisions", width=Inches(5), height=Inches(2.5))

heading(3, "4.1.6 Module de répartition")
paragraph(
    "Ce module permet la répartition automatique des étudiants dans les salles "
    "d'examen en fonction de leur nombre et de la capacité des salles disponibles."
)
add_image_placeholder(doc, "Interface de répartition des salles", width=Inches(5), height=Inches(3))

heading(3, "4.1.7 Module de gestion des modules")
paragraph(
    "Ce module offre une vue d'ensemble des modules avec des statistiques "
    "(nombre d'étudiants inscrits, nombre de professeurs assignés). Il permet "
    "également la gestion des groupes au sein d'un module."
)
add_image_placeholder(doc, "Liste des modules avec statistiques", width=Inches(5), height=Inches(2.5))

heading(3, "4.1.8 Module d'inscription aux examens")
paragraph(
    "Ce module permet d'inscrire les étudiants aux sessions d'examen, de gérer "
    "leur statut et de suivre leurs notes."
)
add_image_placeholder(doc, "Inscription aux examens par module", width=Inches(5), height=Inches(3))

heading(2, "4.2 Interfaces utilisateur")
paragraph(
    "L'interface utilisateur a été conçue pour offrir une expérience fluide et "
    "intuitive. Voici les principales captures d'écran de l'application :"
)

heading(3, "Dashboard")
add_image_placeholder(doc, "Tableau de bord principal", width=Inches(5.5), height=Inches(3))

heading(3, "Navigation et布局")
add_image_placeholder(doc, "Barre de navigation avec menu latéral", width=Inches(5), height=Inches(2.5))

heading(3, "Vue mobile responsive")
add_image_placeholder(doc, "Interface sur écran mobile", width=Inches(3), height=Inches(4))

heading(3, "Interface en arabe (RTL)")
add_image_placeholder(doc, "Interface en langue arabe", width=Inches(5), height=Inches(3))

heading(2, "4.3 Schéma de la base de données")
paragraph(
    "Le schéma relationnel ci-dessous illustre les différentes tables de la base "
    "de données et leurs relations."
)
add_image_placeholder(doc, "Schéma complet de la base de données", width=Inches(5.5), height=Inches(4))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════
heading(1, "Conclusion générale")
paragraph(
    "Ce stage m'a permis de mettre en pratique les compétences acquises durant "
    "ma formation et de les appliquer à un projet concret et fonctionnel. "
    "La plateforme « Prof Space Chariaa » répond aux besoins exprimés par "
    "l'établissement et offre une solution complète pour la gestion de l'espace "
    "professeurs."
)
paragraph(
    "Les objectifs fixés ont été atteints :"
)
bullet("Une plateforme web fonctionnelle et sécurisée")
bullet("Une interface utilisateur intuitive et responsive")
bullet("Des fonctionnalités avancées d'import/export de données")
bullet("Un système de gestion des rôles et des permissions")
bullet("Un支持 multilingue (français/arabe)")
paragraph(
    "Ce projet m'a également permis de développer mes compétences en développement "
    "web full-stack, en particulier avec Laravel et React, et de me familiariser "
    "avec les bonnes pratiques de développement et de gestion de projet."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# BIBLIOGRAPHY
# ══════════════════════════════════════════════════════════════
heading(1, "Bibliographie et webographie")

refs = [
    "Laravel Documentation - https://laravel.com/docs",
    "React Documentation - https://react.dev",
    "Inertia.js Documentation - https://inertiajs.com",
    "Tailwind CSS Documentation - https://tailwindcss.com/docs",
    "MySQL Documentation - https://dev.mysql.com/doc",
    "PhpSpreadsheet - https://phpspreadsheet.readthedocs.io",
    "Laravel Breeze - https://github.com/laravel/breeze",
]
for ref in refs:
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(ref)
    r.font.size = Pt(11)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), "Rapport_de_Stage_Prof_Space_Chariaa.docx")
doc.save(output_path)
print(f"Rapport généré : {output_path}")
