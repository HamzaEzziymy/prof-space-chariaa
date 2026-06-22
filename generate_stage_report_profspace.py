from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "Rapport_de_Stage_ProfSpace_Charia_Fes.docx"


doc = Document()


def setup_section(section):
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


for section in doc.sections:
    setup_section(section)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(12)
styles["Normal"].paragraph_format.line_spacing = 1.25
styles["Normal"].paragraph_format.space_after = Pt(6)

for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    style = styles[style_name]
    style.font.name = "Times New Roman"
    style.font.color.rgb = RGBColor(22, 55, 92)


def shade_cell(cell, fill):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=11, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def p(text="", bold=False, italic=False, align=None):
    para = doc.add_paragraph()
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def h(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(22, 55, 92)
    return heading


def bullets(items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def placeholder(title, note="Insérer ici la capture, le logo, le schéma ou l'information correspondante."):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F8")
    set_cell_text(cell, f"[ {title} ]\n{note}", italic_text := False, size=10, color=(110, 110, 110))
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="B8C2CC"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="B8C2CC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="B8C2CC"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="B8C2CC"/>'
        '</w:tcBorders>'
    )
    tc_pr.append(borders)
    doc.add_paragraph()


def info_table(rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        shade_cell(cells[0], "EAF2F8")
        set_cell_text(cells[0], label, bold=True, size=10, color=(22, 55, 92))
        cells[1].text = value
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


# Cover page
cover = doc.sections[0]
setup_section(cover)

p("ÉCOLE SUPÉRIEURE DE TECHNOLOGIE DE SALÉ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
p("Université Mohammed V de Rabat", align=WD_ALIGN_PARAGRAPH.CENTER)
p("Département : Informatique / Génie Logiciel", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
placeholder("Logo EST Salé", "À remplacer par le logo officiel de l'établissement.")
doc.add_paragraph()
p("RAPPORT DE STAGE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
title = p("Conception et réalisation d'une application web de gestion pédagogique", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
title.runs[0].font.size = Pt(18)
p("Cas de la Faculté de Charia de Fès", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
placeholder("Logo Faculté de Charia de Fès", "À remplacer par le logo officiel de l'organisme d'accueil.")
doc.add_paragraph()

info_table([
    ("Réalisé par", "Nom et prénom : ______________________________"),
    ("Formation", "3ème année - Ingénierie des Applications Web et Mobiles"),
    ("Organisme d'accueil", "Faculté de Charia de Fès"),
    ("Application réalisée", "ProfSpace Chariaa"),
    ("Encadrant pédagogique", "M./Mme ______________________________"),
    ("Encadrant professionnel", "M./Mme ______________________________"),
    ("Période de stage", "Du ____/____/2026 au ____/____/2026"),
    ("Année universitaire", "2025-2026"),
])

p("Soutenu le : ____ / ____ / 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()


h("Remerciements", 1)
p(
    "Je tiens à exprimer mes sincères remerciements à l'ensemble du corps administratif et pédagogique de la Faculté de Charia de Fès pour l'accueil, l'accompagnement et la confiance accordée durant la période de stage. "
    "Je remercie également mon encadrant professionnel pour ses orientations, sa disponibilité et ses remarques constructives tout au long de la réalisation du projet."
)
p(
    "J'adresse mes remerciements à l'École Supérieure de Technologie de Salé, à mes enseignants et à mon encadrant pédagogique pour la formation dispensée durant mon cursus en Ingénierie des Applications Web et Mobiles. "
    "Cette formation m'a permis de mobiliser des compétences en analyse, conception, développement web, gestion de bases de données et amélioration de l'expérience utilisateur."
)
p("Je remercie enfin toutes les personnes ayant contribué, directement ou indirectement, à la réussite de ce stage.")

h("Dédicace", 1)
p("À ma famille, pour son soutien permanent.")
p("À mes enseignants, pour leur accompagnement et leurs conseils.")
p("À toutes les personnes qui m'ont encouragé durant mon parcours académique et professionnel.")
doc.add_page_break()


h("Résumé", 1)
p(
    "Ce rapport présente le travail réalisé dans le cadre d'un stage effectué à la Faculté de Charia de Fès. Le projet porte sur la conception et la réalisation d'une application web de gestion pédagogique intitulée ProfSpace Chariaa. "
    "L'application vise à faciliter la gestion des professeurs, des étudiants, des modules, des inscriptions pédagogiques, des inscriptions aux examens, de la saisie des notes, de la répartition des salles et de la génération de documents."
)
p(
    "La solution développée repose sur une architecture web moderne combinant Laravel pour la partie backend, Inertia.js pour la liaison entre le backend et l'interface, React pour la construction des écrans interactifs, Tailwind CSS pour l'interface utilisateur, ainsi que MySQL ou SQLite selon l'environnement de déploiement. "
    "L'application propose deux espaces principaux : un espace administrateur pour la gestion globale et un espace professeur pour la consultation des modules et la saisie des notes."
)
p(
    "Le travail a permis de mettre en pratique les connaissances acquises en ingénierie des applications web et mobiles, notamment l'analyse des besoins, la modélisation des données, le développement d'interfaces responsives, la sécurisation des accès, l'import/export des données et la génération de rapports."
)
p("Mots-clés : Laravel, React, Inertia.js, gestion pédagogique, notes, examens, administration, professeur.")

h("Abstract", 1)
p(
    "This internship report presents the design and development of ProfSpace Chariaa, a web-based academic management application built for the Faculty of Sharia of Fez. "
    "The application supports administrative and teaching workflows, including teacher management, student records, modules, enrollments, exam registrations, grade entry, room assignment and document generation."
)
p(
    "The system is based on a modern web stack composed of Laravel, Inertia.js, React and Tailwind CSS. It provides role-based access for administrators and professors, bilingual interface support, data import/export features and PDF generation."
)
p("Keywords: Laravel, React, Inertia.js, academic management, grades, exams, web application.")
doc.add_page_break()


h("Table des matières", 1)
p("Insérer ici une table des matières automatique dans Word : Références > Table des matières.")
doc.add_page_break()


h("Liste des figures", 1)
for fig in [
    "Figure 1 : Page d'accueil de l'application",
    "Figure 2 : Tableau de bord administrateur",
    "Figure 3 : Gestion des professeurs",
    "Figure 4 : Gestion des étudiants",
    "Figure 5 : Gestion des modules",
    "Figure 6 : Gestion des inscriptions aux examens",
    "Figure 7 : Tableau de bord professeur",
    "Figure 8 : Saisie des notes par module",
    "Figure 9 : Export PDF du relevé de notes",
]:
    p(fig)
doc.add_page_break()


h("Introduction générale", 1)
p(
    "La transformation numérique occupe aujourd'hui une place essentielle dans les établissements d'enseignement supérieur. Les services administratifs et pédagogiques doivent gérer un volume important d'informations relatives aux étudiants, professeurs, modules, inscriptions, examens et notes. "
    "Lorsque ces informations sont dispersées dans plusieurs fichiers ou traitées manuellement, les risques d'erreur, de perte de temps et de duplication augmentent."
)
p(
    "Dans ce contexte, la Faculté de Charia de Fès avait besoin d'une solution permettant de centraliser et de simplifier les opérations pédagogiques courantes. Le stage a donc consisté à analyser ce besoin, concevoir une solution adaptée, puis développer une application web professionnelle, accessible et évolutive."
)
p(
    "Le présent rapport décrit le contexte du stage, l'organisme d'accueil, la problématique, les objectifs, l'analyse fonctionnelle, la conception technique, la réalisation de l'application, les tests effectués, ainsi que les apports personnels et professionnels de cette expérience."
)

h("Chapitre 1 : Présentation du contexte du stage", 1)
h("1.1 Présentation de l'établissement de formation", 2)
p(
    "L'École Supérieure de Technologie de Salé est un établissement relevant de l'Université Mohammed V de Rabat. Elle assure des formations professionnalisantes dans plusieurs domaines, notamment les sciences et techniques, l'informatique, la gestion et les technologies appliquées. "
    "Dans le cadre de la formation en Ingénierie des Applications Web et Mobiles, les étudiants acquièrent des compétences en développement logiciel, bases de données, architecture web, conception d'interfaces, gestion de projet et déploiement d'applications."
)
p("Informations à compléter ou valider : année de création exacte de la filière, chef de département, responsable de formation, adresse officielle.")
placeholder("Logo EST Salé")

h("1.2 Présentation de l'organisme d'accueil", 2)
p(
    "La Faculté de Charia de Fès est un établissement universitaire public situé à Fès et rattaché à l'Université Sidi Mohamed Ben Abdellah. Elle occupe une place importante dans l'enseignement supérieur marocain lié aux sciences de la Charia, au droit, aux études islamiques et aux domaines associés."
)
p(
    "Le projet réalisé durant le stage s'inscrit dans une logique de modernisation des processus administratifs et pédagogiques, avec pour objectif d'améliorer le suivi des données et de faciliter le travail des administrateurs et des professeurs."
)
p("Informations à compléter : adresse, organigramme officiel, nom du doyen, service d'accueil, encadrant professionnel.")
placeholder("Organigramme de la Faculté de Charia de Fès")

h("1.3 Contexte et besoin général", 2)
p(
    "La gestion pédagogique implique plusieurs opérations répétitives : création des comptes utilisateurs, gestion des professeurs, ajout des étudiants, structuration des filières, niveaux, semestres et modules, inscription des étudiants aux modules, préparation des inscriptions aux examens, saisie des notes et génération de documents. "
    "L'utilisation de fichiers dispersés ou de traitements manuels rend ces opérations plus longues et augmente la probabilité d'erreurs."
)
p(
    "L'application ProfSpace Chariaa a été conçue pour répondre à ce besoin en proposant une plateforme centralisée, sécurisée et adaptée au contexte de la faculté."
)

h("Chapitre 2 : Problématique et objectifs du projet", 1)
h("2.1 Problématique", 2)
p(
    "La problématique principale peut être formulée comme suit : comment concevoir et développer une application web permettant de centraliser, sécuriser et simplifier la gestion pédagogique de la Faculté de Charia de Fès tout en offrant une interface claire aux administrateurs et aux professeurs ?"
)

h("2.2 Objectifs généraux", 2)
bullets([
    "Centraliser les données pédagogiques dans une base unique.",
    "Réduire les traitements manuels liés aux inscriptions, examens et notes.",
    "Permettre aux administrateurs de gérer les structures pédagogiques et les utilisateurs.",
    "Permettre aux professeurs de consulter leurs modules et saisir les notes.",
    "Assurer une interface bilingue français/arabe.",
    "Proposer des exports et documents exploitables par l'administration.",
    "Améliorer la qualité, la rapidité et la traçabilité des opérations.",
])

h("2.3 Objectifs spécifiques", 2)
bullets([
    "Gestion des comptes administrateurs, super administrateurs et professeurs.",
    "Gestion des professeurs, étudiants, filières, niveaux, semestres et modules.",
    "Affectation des professeurs aux modules.",
    "Importation des étudiants et inscriptions depuis des fichiers Excel.",
    "Gestion des inscriptions pédagogiques et des inscriptions aux examens.",
    "Saisie des notes normales, de rattrapage et finales.",
    "Gestion des absences au moyen d'une valeur dédiée.",
    "Génération de relevés ou documents PDF.",
    "Tableaux de bord statistiques pour faciliter le suivi.",
])

h("Chapitre 3 : Analyse fonctionnelle", 1)
h("3.1 Acteurs du système", 2)
info_table([
    ("Super administrateur", "Gère les utilisateurs, les paramètres globaux et dispose des droits les plus élevés."),
    ("Administrateur", "Gère les professeurs, étudiants, modules, inscriptions et examens."),
    ("Professeur", "Accède à son espace, consulte ses modules, saisit/importe/exporte les notes."),
])

h("3.2 Besoins fonctionnels", 2)
bullets([
    "Authentification sécurisée selon le rôle de l'utilisateur.",
    "Tableau de bord administrateur avec indicateurs de suivi.",
    "Tableau de bord professeur avec progression des modules et notes à traiter.",
    "Gestion CRUD des professeurs, étudiants, modules et structures pédagogiques.",
    "Importation et exportation de données au format Excel/CSV.",
    "Gestion des inscriptions pédagogiques par étudiant et par module.",
    "Gestion des inscriptions aux examens avec statut normale/rattrapage/finale.",
    "Saisie et sauvegarde des notes par les professeurs.",
    "Répartition des étudiants dans les salles d'examen.",
    "Génération de documents PDF, notamment des relevés de notes.",
])

h("3.3 Besoins non fonctionnels", 2)
bullets([
    "Sécurité : accès contrôlé par rôles et sessions authentifiées.",
    "Ergonomie : interface claire, responsive et adaptée aux usages administratifs.",
    "Bilinguisme : prise en charge du français et de l'arabe avec sens RTL.",
    "Fiabilité : validation des formulaires et contrôle des données importées.",
    "Maintenabilité : organisation du code selon les conventions Laravel et React.",
    "Évolutivité : architecture permettant l'ajout de nouveaux modules fonctionnels.",
])

h("Chapitre 4 : Conception du système", 1)
h("4.1 Architecture générale", 2)
p(
    "L'application adopte une architecture monolithique moderne : Laravel assure le backend, l'accès aux données, la validation, l'authentification et les règles métier. Inertia.js permet de connecter Laravel aux pages React sans créer une API REST séparée. React assure la construction des interfaces interactives, tandis que Tailwind CSS permet une mise en forme rapide et cohérente."
)
placeholder("Schéma d'architecture générale", "Exemple : Navigateur -> Inertia/React -> Laravel Controllers -> Models/Eloquent -> Base de données.")

h("4.2 Modèle de données", 2)
p("Les principales entités manipulées par l'application sont :")
bullets([
    "User : comptes utilisateurs avec rôle et état d'activation.",
    "Prof : informations liées aux enseignants.",
    "Etudiant : informations personnelles et pédagogiques des étudiants.",
    "Filiere : filières de formation.",
    "Niveau : niveaux d'étude associés aux filières.",
    "Semestre : semestres associés aux niveaux.",
    "Module : unités d'enseignement affectées aux professeurs.",
    "EtudiantModule : inscriptions pédagogiques des étudiants aux modules.",
    "NoteExam : inscriptions aux examens, notes et décisions.",
    "Salle : salles utilisées pour les examens.",
    "AppSetting : paramètres généraux de l'application.",
])
placeholder("Diagramme de classes ou modèle conceptuel de données")

h("4.3 Règles métier importantes", 2)
bullets([
    "Un module peut être affecté à un professeur.",
    "Un étudiant peut être inscrit à plusieurs modules.",
    "Une inscription à un examen possède un statut : normale, rattrapage ou finale.",
    "La note 99 est utilisée pour représenter l'absence.",
    "La décision peut être générée selon la note saisie : validé, non validé ou absent.",
    "Un professeur ne peut gérer que les modules qui lui sont affectés.",
])

h("Chapitre 5 : Technologies utilisées", 1)
info_table([
    ("Laravel 12", "Framework PHP utilisé pour le backend, le routage, les contrôleurs, les modèles et l'authentification."),
    ("React 18", "Bibliothèque JavaScript utilisée pour construire les interfaces utilisateur."),
    ("Inertia.js", "Pont entre Laravel et React, permettant une application fluide sans API REST séparée."),
    ("Tailwind CSS", "Framework CSS utilitaire pour créer une interface responsive et moderne."),
    ("Vite", "Outil de build utilisé pour compiler les assets frontend."),
    ("Ziggy", "Génération des routes Laravel côté JavaScript."),
    ("XLSX", "Traitement des fichiers Excel côté frontend."),
    ("Dompdf / PDF", "Génération de documents PDF."),
    ("Recharts", "Affichage de graphiques statistiques dans le tableau de bord."),
])

h("5.1 Justification des choix techniques", 2)
p(
    "Laravel a été choisi pour sa robustesse, sa structure claire, son ORM Eloquent, son système de migrations et son écosystème adapté aux applications professionnelles. "
    "React permet de créer une interface dynamique et réutilisable à base de composants. Inertia.js simplifie l'intégration entre Laravel et React, évitant la complexité d'une API séparée pour ce type d'application interne. "
    "Tailwind CSS a permis d'accélérer la réalisation d'une interface moderne, responsive et cohérente."
)

h("Chapitre 6 : Réalisation de l'application", 1)
h("6.1 Authentification et gestion des rôles", 2)
p(
    "L'application distingue plusieurs rôles : super administrateur, administrateur et professeur. Chaque rôle dispose d'un espace et de permissions adaptés. "
    "Les administrateurs accèdent aux modules de gestion globale, tandis que les professeurs disposent d'un espace dédié à leurs modules et à la saisie des notes."
)
placeholder("Capture : page de connexion administrateur/professeur")

h("6.2 Tableau de bord administrateur", 2)
p(
    "Le tableau de bord administrateur fournit une vue synthétique sur le système : nombre de professeurs, étudiants, modules, répartition par sexe et statistiques des examens par module. "
    "Il permet à l'administration de disposer rapidement d'indicateurs utiles au suivi pédagogique."
)
placeholder("Capture : tableau de bord administrateur")

h("6.3 Gestion des professeurs", 2)
p(
    "Ce module permet d'ajouter, modifier, consulter et supprimer les informations des professeurs. Il facilite également l'association des comptes utilisateurs aux profils enseignants."
)
placeholder("Capture : gestion des professeurs")

h("6.4 Gestion des étudiants", 2)
p(
    "Le module étudiants permet la gestion des informations personnelles et pédagogiques des étudiants. Il prend en charge l'importation de données et l'affichage détaillé du parcours d'un étudiant."
)
placeholder("Capture : gestion des étudiants")

h("6.5 Gestion de la structure pédagogique", 2)
p(
    "La structure pédagogique regroupe les filières, niveaux, semestres et modules. Cette organisation permet de représenter fidèlement l'architecture pédagogique de la faculté."
)
placeholder("Capture : structure pédagogique")

h("6.6 Inscriptions pédagogiques", 2)
p(
    "L'application permet d'inscrire les étudiants dans les modules correspondants. Elle propose des vues par étudiant et par module, ainsi que des fonctionnalités de recherche et d'importation."
)
placeholder("Capture : inscriptions pédagogiques")

h("6.7 Inscriptions aux examens", 2)
p(
    "Ce module permet de préparer les inscriptions aux examens et d'indiquer le statut de chaque inscription : normale, rattrapage ou finale. Il constitue la base de la saisie des notes par les professeurs."
)
placeholder("Capture : inscriptions aux examens")

h("6.8 Espace professeur", 2)
p(
    "L'espace professeur présente un tableau de bord minimal et professionnel. Il affiche les modules affectés au professeur, la progression des notes saisies et les priorités de traitement. "
    "Le professeur peut ouvrir un module, saisir les notes, importer un fichier, exporter les données ou générer un relevé."
)
placeholder("Capture : tableau de bord professeur")
placeholder("Capture : saisie des notes par module")

h("6.9 Export, import et génération PDF", 2)
p(
    "L'application intègre des mécanismes d'importation et d'exportation afin de réduire le temps de saisie. Les exports peuvent être utilisés pour produire des fichiers de travail ou des documents officiels. "
    "La génération PDF permet de produire des documents tels que des relevés de notes."
)
placeholder("Capture : export PDF ou relevé de notes")

h("Chapitre 7 : Tests et validation", 1)
h("7.1 Stratégie de test", 2)
p("Les tests réalisés ont porté principalement sur les scénarios fonctionnels et la cohérence des données.")
bullets([
    "Test de connexion selon les rôles.",
    "Test de création, modification et suppression des entités principales.",
    "Test d'importation de fichiers Excel/CSV.",
    "Test de validation des notes saisies.",
    "Test du cas d'absence représenté par la note 99.",
    "Test d'accès professeur limité aux modules affectés.",
    "Test d'affichage bilingue français/arabe.",
    "Test de génération des documents PDF.",
])

h("7.2 Exemples de scénarios de validation", 2)
info_table([
    ("Scénario", "Résultat attendu"),
    ("Connexion administrateur", "Accès au tableau de bord administrateur."),
    ("Connexion professeur", "Accès uniquement à l'espace professeur."),
    ("Import étudiants", "Lecture du fichier et signalement des lignes invalides."),
    ("Saisie note normale", "Sauvegarde de la note et mise à jour de la progression."),
    ("Saisie note 99", "Décision affichée comme absent."),
    ("Export PDF", "Génération d'un document exploitable."),
])

h("Chapitre 8 : Difficultés rencontrées et solutions", 1)
info_table([
    ("Difficulté", "Solution apportée"),
    ("Gestion bilingue français/arabe", "Mise en place d'un contexte de langue et adaptation RTL de l'interface."),
    ("Importation de données hétérogènes", "Validation des colonnes, contrôle des erreurs et rapport des lignes rejetées."),
    ("Gestion des statuts d'examen", "Ajout d'un champ statut et adaptation des calculs selon normale/rattrapage/finale."),
    ("Expérience utilisateur", "Création de tableaux de bord, pagination, filtres, badges et messages de retour."),
    ("Génération PDF avec arabe", "Utilisation de polices adaptées et traitement spécifique du texte arabe."),
])

h("Chapitre 9 : Apports du stage", 1)
h("9.1 Apports techniques", 2)
bullets([
    "Renforcement des compétences en Laravel, React, Inertia.js et Tailwind CSS.",
    "Mise en pratique des migrations, relations Eloquent, contrôleurs et middleware.",
    "Conception d'interfaces administratives responsives et bilingues.",
    "Manipulation de fichiers Excel/CSV et génération de documents PDF.",
    "Structuration d'un projet réel avec plusieurs profils utilisateurs.",
])

h("9.2 Apports professionnels", 2)
bullets([
    "Compréhension des besoins d'un établissement universitaire.",
    "Capacité à transformer un besoin métier en solution numérique.",
    "Amélioration de la communication avec les utilisateurs finaux.",
    "Gestion progressive des priorités et des retours.",
    "Sensibilisation à la fiabilité des données pédagogiques.",
])

h("Conclusion générale", 1)
p(
    "Ce stage à la Faculté de Charia de Fès a constitué une expérience importante dans mon parcours de formation en Ingénierie des Applications Web et Mobiles. "
    "Il m'a permis de participer à la conception et au développement d'une application web répondant à un besoin réel de gestion pédagogique."
)
p(
    "L'application ProfSpace Chariaa centralise plusieurs fonctionnalités essentielles : gestion des professeurs, étudiants, modules, inscriptions, examens, notes, tableaux de bord et documents. "
    "Elle permet de réduire la charge manuelle, d'améliorer la lisibilité des données et de faciliter le travail des administrateurs et des professeurs."
)
p(
    "Des perspectives d'amélioration restent possibles, notamment l'ajout de notifications, la mise en place de tests automatisés plus complets, la journalisation avancée des actions, l'intégration d'un système d'archivage par année universitaire et le déploiement dans un environnement de production sécurisé."
)

h("Bibliographie et webographie", 1)
bullets([
    "Université Mohammed V de Rabat - informations générales et rattachement de l'EST Salé : https://www.um5.ac.ma/",
    "École Supérieure de Technologie de Salé - site institutionnel : https://ests.um5.ac.ma/",
    "Université Sidi Mohamed Ben Abdellah de Fès : https://www.usmba.ac.ma/",
    "Faculté de Charia de Fès - site institutionnel : http://sharia.usmba.ac.ma/",
    "Laravel Documentation 12.x : https://laravel.com/docs/12.x",
    "Inertia.js Documentation : https://inertiajs.com/",
    "React Documentation : https://react.dev/learn",
    "Tailwind CSS avec Laravel/Vite : https://tailwindcss.com/docs/installation/framework-guides/laravel/vite",
    "Documentation interne du projet : code source ProfSpace Chariaa, contrôleurs, routes, modèles et interfaces.",
])

h("Annexes", 1)
h("Annexe A : Captures d'écran à insérer", 2)
for item in [
    "Page d'accueil",
    "Connexion administrateur",
    "Tableau de bord administrateur",
    "Gestion des professeurs",
    "Gestion des étudiants",
    "Gestion des modules",
    "Inscriptions pédagogiques",
    "Inscriptions aux examens",
    "Tableau de bord professeur",
    "Saisie des notes",
    "Export PDF",
]:
    placeholder(item)

h("Annexe B : Informations à compléter", 2)
bullets([
    "Nom et prénom de l'étudiant.",
    "Nom de l'encadrant pédagogique.",
    "Nom de l'encadrant professionnel.",
    "Période exacte du stage.",
    "Adresse complète de la Faculté de Charia de Fès.",
    "Logo officiel de l'EST Salé.",
    "Logo officiel de la Faculté de Charia de Fès.",
    "Captures d'écran finales de l'application.",
    "Éventuelle validation de l'organigramme institutionnel.",
])

doc.save(OUTPUT)
print(OUTPUT)
