# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path("Rapport_de_Stage_ProfSpace_30_Pages.docx")


def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, fill=None):
    if fill:
        shade_cell(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, 10, bold=bold)


def configure(document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, (31, 78, 121)),
        ("Heading 2", 14, (47, 117, 181)),
        ("Heading 3", 12, (70, 70, 70)),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)


def paragraph(document, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, size=11, bold=False, italic=False):
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r)


def numbered(document, text):
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r)


def heading(document, text, level=1):
    document.add_heading(text, level=level)


def page_title(document, text):
    heading(document, text, 1)


def placeholder_box(document, title, lines=3):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F2F6FA")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 10, bold=True, italic=True, color=(90, 90, 90))
    for _ in range(lines):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("[espace reserve]")
        set_font(r, 9, italic=True, color=(120, 120, 120))
    document.add_paragraph()


def cover(document):
    placeholder_box(document, "[Logo EST Sale]        [Logo Faculte Charia de Fes]        [Logo Universite]", 2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rapport de stage")
    set_font(r, 22, bold=True, color=(31, 78, 121))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Application ProfSpace : plateforme de gestion pedagogique et de suivi des examens")
    set_font(r, 15, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Stage effectue a la Faculte Charia de Fes")
    set_font(r, 13, italic=True, color=(80, 80, 80))

    document.add_paragraph()

    table = document.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    data = [
        ("Realise par", "[Nom et prenom du stagiaire]"),
        ("Formation", "3eme annee - Ingenierie des Applications Web et Mobiles"),
        ("Etablissement d'origine", "Ecole Superieure de Technologie de Sale"),
        ("Organisme d'accueil", "Faculte Charia de Fes"),
        ("Encadrant pedagogique", "[Nom de l'encadrant pedagogique]"),
        ("Encadrant professionnel", "[Nom de l'encadrant professionnel]"),
        ("Periode de stage", "[Date de debut] - [Date de fin]"),
        ("Annee universitaire", "[2025-2026]"),
    ]
    for row, (label, value) in zip(table.rows, data):
        set_cell(row.cells[0], label, bold=True, fill="D9EAF7")
        set_cell(row.cells[1], value)

    document.add_paragraph()
    paragraph(document, "Ce document constitue une version complete et editable du rapport de stage. Les noms, signatures, logos, captures d'ecran et informations non confirmees sont laisses sous forme d'espaces reserves.", WD_ALIGN_PARAGRAPH.CENTER, italic=True)
    document.add_page_break()


def add_preliminary_pages(document):
    page_title(document, "Dedicace")
    paragraph(document, "Je dedie ce travail a ma famille, pour son soutien constant, sa patience et ses encouragements durant mon parcours de formation. Cette experience de stage represente une etape importante dans mon apprentissage, et elle n'aurait pas eu la meme valeur sans l'accompagnement moral des personnes qui m'ont encourage a progresser.")
    paragraph(document, "Je dedie egalement ce rapport a mes enseignants de l'Ecole Superieure de Technologie de Sale, qui ont contribue a construire les bases scientifiques, techniques et methodologiques necessaires a la realisation de ce projet.")
    paragraph(document, "Enfin, je dedie ce travail a toutes les personnes qui participent, directement ou indirectement, a l'amelioration du fonctionnement administratif et pedagogique des etablissements universitaires. La transformation numerique n'a de sens que lorsqu'elle facilite le travail humain et rend les services plus accessibles.")
    document.add_page_break()

    page_title(document, "Remerciements")
    paragraph(document, "Avant de presenter le contenu technique du projet, je tiens a exprimer mes sinceres remerciements a toutes les personnes qui ont contribue a la reussite de ce stage. Je remercie en premier lieu la Faculte Charia de Fes pour m'avoir accueilli et permis de realiser un projet concret, directement lie aux besoins de gestion pedagogique.")
    paragraph(document, "Je remercie mon encadrant professionnel, [Nom de l'encadrant professionnel], pour ses orientations, ses remarques et sa disponibilite. Ses retours ont permis d'adapter progressivement l'application au contexte reel d'utilisation et de mieux comprendre les contraintes du travail administratif.")
    paragraph(document, "Je remercie egalement mon encadrant pedagogique, [Nom de l'encadrant pedagogique], pour son suivi et ses conseils. Je remercie enfin l'ensemble du corps professoral de l'EST Sale pour la qualite de la formation dispensee, notamment en developpement web, bases de donnees, conception logicielle, ergonomie et gestion de projet.")
    paragraph(document, "Ce stage a ete une occasion precieuse d'appliquer des connaissances techniques dans un environnement reel, avec des utilisateurs, des contraintes, des donnees sensibles et des objectifs institutionnels.")
    document.add_page_break()

    page_title(document, "Resume")
    paragraph(document, "Le present rapport de stage decrit la conception et la realisation de ProfSpace, une application web destinee a faciliter la gestion pedagogique et le suivi des examens au sein de la Faculte Charia de Fes. Le projet a ete realise dans le cadre de la formation de 3eme annee en Ingenierie des Applications Web et Mobiles a l'EST Sale.")
    paragraph(document, "L'application permet de centraliser plusieurs operations : gestion des etudiants, gestion des professeurs, gestion des modules, structure pedagogique, inscriptions pedagogiques, inscriptions aux examens, saisie des notes, import/export Excel, generation de releves PDF et visualisation de statistiques.")
    paragraph(document, "Le projet s'appuie sur Laravel pour le backend, Inertia.js pour la communication entre Laravel et React, React pour les interfaces interactives, Tailwind CSS pour le design responsive, Recharts pour les graphiques et des bibliotheques PDF/Excel pour les documents. Une attention particuliere a ete accordee a l'experience utilisateur, au bilinguisme francais/arabe, au mode RTL et a la reduction de la charge administrative.")
    paragraph(document, "Mots-cles : Laravel, React, Inertia.js, gestion pedagogique, notes, examens, UI/UX, bilinguisme, RTL, Excel, PDF.")

    heading(document, "Abstract", 2)
    paragraph(document, "This internship report presents the design and development of ProfSpace, a web application created to support academic management and examination tracking at the Faculty of Sharia in Fez. The project was carried out as part of the third-year Web and Mobile Application Engineering program at EST Sale.")
    paragraph(document, "The application centralizes student management, professor management, modules, academic structure, enrollments, exam registrations, grade entry, Excel import/export, PDF generation and statistical dashboards. The technical stack is based on Laravel, Inertia.js, React, Tailwind CSS, Recharts and document-processing libraries.")
    document.add_page_break()

    page_title(document, "Table des matieres")
    paragraph(document, "La table des matieres doit etre generee automatiquement dans Microsoft Word apres l'ouverture du document : References > Table des matieres > Table automatique.", italic=True)
    for item in [
        "Introduction generale",
        "Chapitre 1 : Presentation du contexte et de l'organisme d'accueil",
        "Chapitre 2 : Analyse de l'existant et problematique",
        "Chapitre 3 : Analyse fonctionnelle et conception",
        "Chapitre 4 : Realisation technique de l'application",
        "Chapitre 5 : Qualite UI/UX et approche humaine",
        "Chapitre 6 : Tests, validation et limites",
        "Chapitre 7 : Bilan personnel et perspectives",
        "Conclusion generale",
        "Bibliographie et webographie",
        "Annexes",
    ]:
        bullet(document, item)
    document.add_page_break()

    page_title(document, "Liste des figures, tableaux et abreviations")
    heading(document, "Liste des figures", 2)
    for fig in [
        "Figure 1 : Logo de l'organisme d'accueil [a inserer]",
        "Figure 2 : Architecture generale de l'application [a inserer]",
        "Figure 3 : Modele conceptuel de donnees [a inserer]",
        "Figure 4 : Tableau de bord administrateur [capture a inserer]",
        "Figure 5 : Tableau de bord professeur [capture a inserer]",
        "Figure 6 : Page de saisie des notes [capture a inserer]",
        "Figure 7 : Graphique des statistiques des examens par module [capture a inserer]",
    ]:
        bullet(document, fig)
    heading(document, "Abreviations", 2)
    table = document.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    entries = [
        ("API", "Application Programming Interface"),
        ("CRUD", "Create, Read, Update, Delete"),
        ("CNE", "Code National de l'Etudiant"),
        ("MVC", "Model View Controller"),
        ("PDF", "Portable Document Format"),
        ("RTL", "Right To Left"),
        ("UI/UX", "User Interface / User Experience"),
    ]
    for row, (abbr, meaning) in zip(table.rows, entries):
        set_cell(row.cells[0], abbr, bold=True, fill="EAF2F8")
        set_cell(row.cells[1], meaning)
    document.add_page_break()


REPORT_PAGES = [
    {
        "title": "Introduction generale",
        "paragraphs": [
            "La transformation numerique occupe aujourd'hui une place importante dans les etablissements d'enseignement superieur. Les services administratifs et pedagogiques manipulent un volume important de donnees : listes d'etudiants, modules, professeurs, inscriptions, examens, notes et decisions. Lorsque ces informations sont reparties entre plusieurs fichiers ou traitements manuels, la gestion devient plus lente et plus vulnerable aux erreurs.",
            "Dans ce contexte, le projet ProfSpace a ete concu comme une solution web permettant de centraliser les donnees pedagogiques et de simplifier le suivi des examens. L'application cherche a rapprocher l'administration et les professeurs dans un meme environnement numerique, tout en respectant les responsabilites de chaque profil utilisateur.",
            "Ce rapport presente les differentes etapes de realisation du projet, depuis l'analyse du besoin jusqu'a la conception, la realisation technique, l'amelioration de l'interface, les tests et les perspectives. Une attention particuliere est accordee a la qualite de l'interface utilisateur, a l'experience d'utilisation et a l'approche humaine du projet.",
            "L'objectif n'etait pas uniquement de produire une application fonctionnelle. Il s'agissait aussi de concevoir un outil comprehensible, rassurant et adapte a des utilisateurs qui n'ont pas tous le meme niveau de familiarite avec les outils informatiques.",
        ],
    },
    {
        "title": "Chapitre 1 : Presentation du contexte du stage",
        "paragraphs": [
            "Le stage a ete effectue dans le cadre de la formation de 3eme annee en Ingenierie des Applications Web et Mobiles a l'Ecole Superieure de Technologie de Sale. Cette formation combine des competences en programmation, developpement web, developpement mobile, conception de bases de donnees, genie logiciel, securite, interfaces utilisateur et gestion de projet.",
            "Le choix d'un projet de gestion pedagogique est coherent avec cette formation, car il mobilise plusieurs competences transversales. Il ne s'agit pas seulement de creer des pages web, mais de concevoir une application complete avec authentification, roles, formulaires, tableaux, import/export, generation de documents et tableaux de bord.",
            "L'organisme d'accueil, la Faculte Charia de Fes, represente un environnement reel ou les donnees pedagogiques ont une valeur importante. Une erreur dans une inscription, une note ou une association de module peut entrainer des consequences administratives. Le projet devait donc etre pense avec rigueur et prudence.",
            "Le stage a ainsi permis de passer d'une logique d'exercice academique a une logique de produit utile. Cette transition est essentielle pour un futur ingenieur d'applications web et mobiles, car elle oblige a prendre en compte les utilisateurs, les contraintes institutionnelles et les workflows reels.",
        ],
        "placeholder": "[Inserez ici une image de l'etablissement d'accueil ou son logo]",
    },
    {
        "title": "Presentation de l'organisme d'accueil",
        "paragraphs": [
            "La Faculte Charia de Fes est un etablissement universitaire marocain rattache a l'Universite Sidi Mohamed Ben Abdellah. Elle intervient dans un contexte d'enseignement superieur ou la gestion des parcours, des modules et des examens necessite une organisation precise.",
            "Dans le cadre de ce rapport, certaines informations institutionnelles comme les noms des responsables, l'organigramme detaille ou les dates exactes doivent etre completees et validees avec l'organisme d'accueil. Les espaces reserves permettent d'ajouter ces informations sans modifier la structure generale du document.",
            "L'environnement d'accueil a oriente les choix du projet. La presence du francais et de l'arabe dans les donnees, la necessite d'utiliser des fichiers Excel et la sensibilite des notes ont pousse vers une application bilingue, securisee et capable de gerer des imports/exports.",
            "Le projet s'inscrit donc dans une logique de modernisation progressive. Il respecte les pratiques existantes, comme l'usage des feuilles Excel, tout en proposant une centralisation plus fiable et une interface plus claire.",
        ],
        "placeholder": "[Organigramme ou presentation institutionnelle a inserer]",
    },
    {
        "title": "Contexte de la formation et competences mobilisees",
        "paragraphs": [
            "La formation en Ingenierie des Applications Web et Mobiles prepare les etudiants a concevoir des solutions logicielles completes. Dans ce projet, plusieurs blocs de competences ont ete mobilises : analyse fonctionnelle, modelisation, developpement backend, developpement frontend, design responsive, securite applicative, manipulation de fichiers et documentation.",
            "Le backend Laravel a permis de mettre en place les routes, controleurs, validations, modeles Eloquent et middlewares. Le frontend React a permis de construire des interfaces interactives adaptees aux besoins des utilisateurs. Inertia.js a joue le role de liaison entre les deux couches, en evitant la complexite d'une API separee.",
            "Les competences UI/UX ont egalement ete importantes. Une application administrative doit rester dense, mais lisible. Il fallait donc choisir les bons composants : cartes statistiques, tableaux pagines, modales, filtres, boutons d'action, graphiques et messages de retour.",
            "Ce stage a aussi necessite des competences humaines : ecoute, comprehension du contexte, adaptation aux retours, patience dans les ajustements et capacite a transformer un besoin vague en fonctionnalites concretes.",
        ],
    },
    {
        "title": "Chapitre 2 : Analyse de l'existant",
        "paragraphs": [
            "L'analyse de l'existant consiste a comprendre comment les taches etaient ou pourraient etre realisees avant la mise en place de l'application. Dans un contexte universitaire, beaucoup d'operations reposent souvent sur des fichiers Excel, des echanges par email ou des documents imprimes.",
            "Ces outils sont utiles et familiers, mais ils presentent des limites lorsque le volume de donnees augmente. Les risques les plus frequents sont les doublons, les oublis, les erreurs de copie, les fichiers obsoletes, les versions concurrentes et le manque de visibilite globale.",
            "La gestion des notes est un exemple critique. Lorsqu'un professeur saisit des notes dans un fichier separe, l'administration doit ensuite les recuperer, les verifier, les consolider et parfois corriger les erreurs. Ce processus prend du temps et peut creer de l'incertitude.",
            "Le projet ProfSpace repond a cette situation par une centralisation progressive : les donnees sont stockees dans une base, les utilisateurs travaillent selon leurs roles, les imports restent possibles et les statistiques donnent une vision immediate de l'avancement.",
        ],
    },
    {
        "title": "Problematique du projet",
        "paragraphs": [
            "La problematique principale peut etre formulee ainsi : comment concevoir une application web permettant de simplifier la gestion pedagogique et le suivi des examens, tout en garantissant la fiabilite des donnees, la securite des acces et une experience utilisateur adaptee aux administrateurs et aux professeurs ?",
            "Cette problematique contient plusieurs dimensions. La premiere est fonctionnelle : l'application doit couvrir les operations essentielles, depuis la structure pedagogique jusqu'a la saisie des notes. La deuxieme est technique : le systeme doit etre organise, maintenable et capable de traiter des fichiers. La troisieme est humaine : l'application doit etre facile a comprendre et ne pas compliquer le travail quotidien.",
            "Le projet doit aussi prendre en compte le bilinguisme. Dans le contexte marocain, les informations peuvent exister en francais et en arabe. Il ne suffit donc pas de traduire quelques labels ; il faut aussi prevoir les noms arabes, le sens de lecture RTL et l'affichage correct dans les documents PDF.",
            "Enfin, l'application doit inspirer confiance. Les utilisateurs doivent savoir ce qui a ete sauvegarde, ce qui reste en attente, quelles lignes d'import ont echoue et quelles actions sont autorisees.",
        ],
    },
    {
        "title": "Objectifs generaux et specifiques",
        "paragraphs": [
            "L'objectif general du projet est de developper une application web centralisee pour la gestion pedagogique et le suivi des examens. Cet objectif se decline en plusieurs objectifs specifiques, chacun correspondant a un besoin utilisateur clairement identifiable.",
            "Pour l'administration, l'application doit permettre de gerer les professeurs, les etudiants, les modules, la structure pedagogique, les inscriptions pedagogiques et les inscriptions aux examens. Elle doit aussi fournir des statistiques et des outils d'import/export.",
            "Pour les professeurs, l'application doit proposer un espace simple et cible : voir les modules affectes, identifier les notes en attente, saisir les notes, importer un fichier, exporter une liste et generer un releve PDF.",
            "Pour le systeme global, les objectifs sont la securite des roles, la reduction des erreurs, la traçabilite, le support bilingue et la possibilite d'evoluer vers d'autres fonctionnalites.",
        ],
        "bullets": [
            "Centraliser les donnees pedagogiques.",
            "Reduire la manipulation manuelle des fichiers.",
            "Faciliter la saisie des notes par les professeurs.",
            "Produire des statistiques d'avancement.",
            "Respecter le contexte francais/arabe.",
        ],
    },
    {
        "title": "Cahier des charges fonctionnel",
        "paragraphs": [
            "Le cahier des charges fonctionnel a ete construit autour des principaux flux de l'application. Chaque flux correspond a une serie d'actions attendues par un profil utilisateur.",
            "L'administrateur doit pouvoir creer, modifier, supprimer, rechercher, filtrer, importer et exporter les entites pedagogiques. Ces actions concernent notamment les etudiants, les professeurs, les modules, les filieres, les niveaux et les semestres.",
            "Le professeur doit pouvoir consulter uniquement les modules qui lui sont associes. Cette limitation simplifie son interface et evite l'acces a des donnees qui ne le concernent pas. Il doit ensuite pouvoir saisir les notes des etudiants inscrits a l'examen.",
            "Le super administrateur dispose de fonctionnalites plus sensibles : gestion des utilisateurs, activation/desactivation des comptes, parametrage global et maintenance. Cette separation correspond au principe du moindre privilege.",
        ],
        "bullets": [
            "Authentification et gestion des roles.",
            "Gestion de la structure pedagogique.",
            "Gestion des etudiants et professeurs.",
            "Gestion des modules et affectations.",
            "Inscriptions pedagogiques et inscriptions aux examens.",
            "Saisie, import, export et PDF des notes.",
        ],
    },
    {
        "title": "Cahier des charges non fonctionnel",
        "paragraphs": [
            "Les exigences non fonctionnelles sont aussi importantes que les exigences fonctionnelles. Elles definissent la qualite attendue de l'application : securite, performance, ergonomie, maintenabilite, compatibilite et accessibilite.",
            "La securite est essentielle car l'application manipule des donnees personnelles et des notes. Les routes doivent etre protegees, les roles doivent etre verifies et les donnees envoyees par l'utilisateur doivent etre validees cote serveur.",
            "L'ergonomie est egalement prioritaire. Les utilisateurs doivent pouvoir realiser leurs taches rapidement. L'interface doit eviter les longues pages confuses, proposer des recherches et filtres, donner des retours clairs et presenter les informations selon une hierarchie visuelle logique.",
            "La maintenabilite repose sur une architecture claire. Le projet s'appuie sur Laravel, ses controleurs, ses modeles Eloquent, ses middlewares et Inertia.js pour garder une organisation lisible. Cette structure facilite les evolutions futures.",
        ],
        "bullets": [
            "Securite des acces et validation serveur.",
            "Interface responsive et lisible.",
            "Support du mode sombre.",
            "Support francais/arabe et RTL.",
            "Gestion des erreurs d'import.",
            "Code organise et extensible.",
        ],
    },
    {
        "title": "Chapitre 3 : Methodologie adoptee",
        "paragraphs": [
            "La realisation du projet a suivi une approche iterative. Plutot que de chercher a produire toutes les fonctionnalites en une seule fois, le travail a ete organise par modules. Cette methode permet de construire progressivement l'application et d'ameliorer l'interface apres chaque retour.",
            "Une premiere phase a consiste a identifier les donnees centrales : utilisateurs, professeurs, etudiants, filieres, niveaux, semestres, modules, inscriptions et notes. Ensuite, les routes et controleurs ont ete mis en place pour exposer les operations principales.",
            "La phase suivante a concerne l'interface. Les pages React ont ete construites autour de composants reutilisables : layouts, cartes, tableaux, modales, graphiques et formulaires. Les ajustements UI/UX ont ete realises au fur et a mesure, notamment pour le dashboard professeur, le graphique des examens et le mode arabe.",
            "Cette methode iterative est adaptee a un stage, car elle permet de montrer rapidement des resultats, de recevoir des remarques et d'ameliorer l'application sans attendre la fin du projet.",
        ],
    },
    {
        "title": "Architecture generale de l'application",
        "paragraphs": [
            "L'application repose sur une architecture Laravel/Inertia/React. Laravel gere le backend, les routes, les controleurs, les modeles, les validations et les middlewares. React gere l'interface utilisateur. Inertia.js assure la communication entre les deux sans imposer la creation d'une API REST separee.",
            "Cette architecture est particulierement adaptee a une application de gestion. Elle permet de conserver la puissance de Laravel cote serveur tout en offrant une interface moderne et reactive cote client. Les pages React recoivent les donnees preparees par les controleurs Laravel.",
            "Le routage est separe en deux espaces : les routes web principales pour l'administration, et les routes professeur dans `routes/prof.php`. Cette separation rend l'application plus lisible et facilite la maintenance.",
            "Les donnees communes, comme l'utilisateur connecte, les parametres de l'application et les modules du professeur, sont partagees via le middleware Inertia. Cela evite de repeter des requetes dans chaque page.",
        ],
        "placeholder": "[Schema d'architecture generale a inserer]",
    },
    {
        "title": "Modelisation des donnees",
        "paragraphs": [
            "La modelisation des donnees constitue le coeur du projet. Elle doit representer correctement la realite pedagogique. Le modele `User` gere les comptes et les roles. Le modele `Prof` represente la fiche professeur et se rattache a un utilisateur. Le modele `Etudiant` stocke les informations administratives et academiques des etudiants.",
            "La structure pedagogique est organisee avec les filieres, niveaux, semestres et modules. Un module appartient a un semestre, un semestre appartient a un niveau, et un niveau appartient a une filiere. Cette hierarchie permet de classer les modules et les etudiants de facon coherente.",
            "L'inscription pedagogique est representee par `EtudiantModule`, qui relie un etudiant a un module. L'inscription a l'examen et les notes sont representees par `NoteExam`. Cette separation permet de distinguer le fait qu'un etudiant suit un module et le fait qu'il est inscrit a un examen avec des notes.",
            "Le modele `NoteExam` gere les notes de session normale, rattrapage et finale. Il calcule aussi les decisions en francais et en arabe, ce qui permet d'afficher ou d'exporter les resultats plus facilement.",
        ],
        "placeholder": "[Modele conceptuel ou diagramme de classes a inserer]",
    },
    {
        "title": "Securite, authentification et roles",
        "paragraphs": [
            "L'application distingue plusieurs profils : super administrateur, administrateur et professeur. Cette separation est appliquee par des middlewares Laravel. Le middleware admin limite l'acces aux administrateurs, le middleware super admin limite les fonctions sensibles et le middleware prof protege l'espace professeur.",
            "L'authentification est basee sur les mecanismes Laravel. Les utilisateurs disposent d'un compte avec email, mot de passe, role, statut actif et photo de profil. Le projet prevoit aussi l'obligation de changer le mot de passe, ce qui renforce la securite lors de la premiere connexion ou apres creation d'un compte.",
            "La securite ne concerne pas seulement l'acces aux pages. Elle concerne aussi les actions : creation, modification, suppression, importation et sauvegarde de notes. Les donnees doivent etre validees cote serveur, meme si l'interface effectue deja des controles.",
            "Une amelioration recommandee consiste a renforcer la verification de chaque ligne de note lors de la sauvegarde : chaque identifiant d'inscription envoye par le professeur doit etre confirme comme appartenant au module et au professeur connecte.",
        ],
    },
    {
        "title": "Chapitre 4 : Tableau de bord administrateur",
        "paragraphs": [
            "Le tableau de bord administrateur offre une vision globale de l'application. Il affiche des indicateurs comme le nombre de professeurs, d'etudiants, de modules, de modules actifs, d'utilisateurs actifs et de notes saisies. Il presente aussi des graphiques et des listes recentes.",
            "L'objectif du tableau de bord est d'aider l'administration a comprendre rapidement l'etat du systeme. Au lieu de parcourir plusieurs pages, l'utilisateur dispose d'une synthese visuelle. Les cartes statistiques servent de resume, tandis que les graphiques permettent de comparer les donnees.",
            "Le graphique des examens par module est un element important. Il affiche le pourcentage de notes saisies pour la session normale et le rattrapage. Les couleurs vert et orange permettent de distinguer les deux modes. L'axe vertical de 0 a 100% facilite la comparaison entre les modules.",
            "Ce dashboard illustre une approche UX orientee decision : l'utilisateur ne voit pas seulement des chiffres, il voit des informations qui peuvent l'aider a agir, par exemple identifier un module qui manque encore de notes.",
        ],
        "placeholder": "[Capture du dashboard administrateur a inserer]",
    },
    {
        "title": "Gestion des etudiants",
        "paragraphs": [
            "La gestion des etudiants est l'un des modules les plus importants de l'application. Elle permet de consulter, rechercher, filtrer, ajouter, modifier, supprimer, importer et exporter les donnees des etudiants. Les informations gerees incluent notamment CNE, CIN, nom, prenom, sexe, email, telephone, niveau et photo.",
            "Le projet prend en charge l'import de fichiers Excel ou CSV. Cette fonctionnalite est essentielle dans un contexte administratif car les listes d'etudiants existent souvent sous forme de fichiers. L'application detecte les lignes invalides, les doublons et les champs manquants.",
            "L'interface de gestion des etudiants doit rester claire malgre la densite des donnees. Les filtres par sexe, filiere et niveau permettent de reduire rapidement la liste. La pagination evite d'afficher un volume trop important en une seule page.",
            "La fiche etudiant permet de consulter les informations detaillees et les modules inscrits. Cette vue detaillee est utile pour verifier rapidement la situation d'un etudiant sans manipuler plusieurs fichiers.",
        ],
        "placeholder": "[Capture de la liste ou fiche etudiant a inserer]",
    },
    {
        "title": "Gestion des professeurs et des modules",
        "paragraphs": [
            "La gestion des professeurs repose sur deux entites : le compte utilisateur et la fiche professeur. Cette separation permet de gerer a la fois l'authentification et les informations professionnelles comme le grade, le CIN et le telephone.",
            "Les modules sont definis par leur nom en francais, nom en arabe, code, coefficient, type, semestre et professeur responsable. L'assignation d'un professeur a un module est une operation centrale, car elle determine ce que le professeur verra dans son espace personnel.",
            "La gestion des modules propose des filtres, une recherche, une pagination, un import Excel/CSV et un export personnalise. Ces fonctionnalites reduisent le travail manuel et facilitent la maintenance de la structure pedagogique.",
            "Du point de vue UI/UX, les pages de gestion utilisent des tableaux, des modales et des boutons d'action. Cette approche est adaptee a une application administrative, ou l'utilisateur doit pouvoir comparer, modifier et controler les donnees rapidement.",
        ],
    },
    {
        "title": "Structure pedagogique",
        "paragraphs": [
            "La structure pedagogique organise les filieres, niveaux et semestres. Elle constitue la base sur laquelle reposent les modules et les inscriptions. Une structure claire evite les incoherences dans les affectations et facilite la navigation dans les donnees.",
            "L'application permet de gerer les filieres, les niveaux et les semestres avec des noms en francais et en arabe. Cette double representation est importante pour respecter le contexte bilingue de l'etablissement.",
            "Le choix de regrouper la structure pedagogique dans un espace dedie simplifie l'administration. Au lieu de disperser les niveaux et semestres dans plusieurs menus, l'utilisateur dispose d'une vue plus coherente.",
            "Cette partie peut etre enrichie dans le futur par une visualisation arborescente plus avancee, montrant les filieres, les niveaux, les semestres et les modules dans une meme representation.",
        ],
        "placeholder": "[Capture de la structure pedagogique a inserer]",
    },
    {
        "title": "Inscriptions pedagogiques",
        "paragraphs": [
            "Les inscriptions pedagogiques relient les etudiants aux modules. Cette relation est indispensable car elle definit quels etudiants sont concernes par chaque module. Elle sert ensuite de base pour creer les inscriptions aux examens.",
            "L'application propose deux modes de consultation : par module et par etudiant. Ce choix est important pour l'ergonomie. L'administration peut chercher les etudiants d'un module ou les modules d'un etudiant selon le besoin du moment.",
            "L'import d'inscriptions permet de traiter de grands volumes. Le projet prend egalement en charge un format horizontal matriciel, ou chaque ligne represente un etudiant et chaque colonne un module. Cette approche est proche de certaines pratiques Excel existantes.",
            "Les rapports d'erreurs sont essentiels. Ils permettent de comprendre si un CNE est introuvable, si un code module n'existe pas ou si une valeur est invalide. L'utilisateur peut ainsi corriger son fichier sans tout recommencer.",
        ],
    },
    {
        "title": "Inscriptions aux examens",
        "paragraphs": [
            "Les inscriptions aux examens sont gerees par le modele `NoteExam`. Cette entite associe une inscription pedagogique a un examen, avec un numero d'examen, un statut et des notes. Les statuts principaux sont normale, rattrapage et finale.",
            "Cette partie est sensible car elle fait le lien entre les donnees administratives et les notes. Elle doit garantir que seuls les etudiants concernes apparaissent dans la liste de saisie. Elle doit aussi eviter de proposer en rattrapage un etudiant deja valide en session normale.",
            "L'administration peut creer les inscriptions, modifier les statuts, appliquer un statut par lot, importer des donnees et exporter les inscriptions. Ces fonctionnalites permettent de preparer le travail des professeurs.",
            "Le choix de centraliser les inscriptions aux examens dans l'application permet de suivre l'avancement des notes et de produire des statistiques fiables.",
        ],
    },
    {
        "title": "Espace professeur",
        "paragraphs": [
            "L'espace professeur est volontairement separe de l'espace admin. Un enseignant n'a pas besoin de voir toute la gestion pedagogique ; il a surtout besoin de voir ses modules, les etudiants concernes et les notes a saisir.",
            "Le tableau de bord professeur affiche une carte de bienvenue, des statistiques de synthese, une priorite actuelle et la progression des modules. Cette interface est orientee action : elle indique au professeur ce qu'il doit traiter en premier.",
            "La carte de priorite actuelle est une decision UX importante. Elle reduit la recherche manuelle et aide le professeur a reprendre son travail rapidement. La carte de progression permet de suivre plusieurs modules sans surcharger l'ecran.",
            "Lorsque le professeur a beaucoup de modules, l'interface doit rester compacte. La pagination interne ou les boutons de navigation entre modules permettent de garder une experience propre et lisible.",
        ],
        "placeholder": "[Capture du tableau de bord professeur a inserer]",
    },
    {
        "title": "Saisie des notes",
        "paragraphs": [
            "La saisie des notes est l'une des fonctionnalites les plus critiques du projet. Elle concerne directement les resultats des etudiants et doit donc etre rapide, claire et fiable. L'application affiche les etudiants inscrits au module et propose un champ de note pour chacun.",
            "Le systeme accepte les notes entre 0 et 20, ainsi que la valeur 99 pour representer l'absence. La decision est calculee en fonction de la note : valide, non valide ou absent. Ce retour immediat aide le professeur a reperer les erreurs.",
            "L'interface prend en charge la navigation clavier avec la touche Enter pour passer a la ligne suivante. Ce detail est important car la saisie des notes est repetitive. Une bonne UX doit permettre de travailler vite sans fatigue excessive.",
            "L'application affiche aussi un indicateur de modifications non sauvegardees. Ce retour visuel limite le risque de quitter la page sans enregistrer. Les messages de succes ou d'erreur renforcent la confiance dans le systeme.",
        ],
        "placeholder": "[Capture de la page de saisie des notes a inserer]",
    },
    {
        "title": "Import, export et generation PDF",
        "paragraphs": [
            "Le projet integre fortement les fichiers Excel et CSV. Cette decision est pragmatique : les services administratifs et les professeurs utilisent souvent Excel pour preparer ou verifier des listes. L'application ne rejette pas cette habitude, elle l'encadre.",
            "Les imports verifient les colonnes obligatoires, les formats, les doublons et les lignes invalides. Les exports permettent de recuperer les donnees dans un format exploitable. Les rapports d'erreurs aident a corriger les fichiers sources.",
            "La generation PDF permet de produire des releves de notes. Le projet tient compte de la langue et utilise une police adaptee pour l'arabe. Cette attention est importante car un document imprime ou partage doit rester lisible et professionnel.",
            "Cette combinaison import/export/PDF montre que l'application n'est pas isolee. Elle s'integre dans un ecosysteme de travail ou les documents restent necessaires.",
        ],
        "placeholder": "[Exemple de releve PDF ou modal d'import a inserer]",
    },
    {
        "title": "Chapitre 5 : Qualite de l'interface utilisateur",
        "paragraphs": [
            "La qualite de l'interface utilisateur est un facteur essentiel pour l'adoption du projet. Une application peut etre techniquement fonctionnelle mais peu utilisee si elle est confuse, lente ou intimidante. Dans ProfSpace, l'interface a ete pensee pour rester claire et professionnelle.",
            "Le design utilise des cartes, des tableaux, des modales, des icones, des badges et des graphiques. Les couleurs sont utilisees pour transmettre du sens : le vert pour la progression ou la validation, l'orange pour l'attention ou le rattrapage, le rouge pour les erreurs.",
            "La navigation laterale permet d'acceder rapidement aux modules de l'application. Elle peut etre repliee, ce qui ameliore l'espace disponible. Le mode sombre ajoute un confort visuel pour certains utilisateurs.",
            "La hierarchie visuelle est importante. Les informations principales apparaissent sous forme de cartes ou de titres, tandis que les details sont presentes dans des tableaux. Cette organisation evite a l'utilisateur de se perdre dans une masse de donnees.",
        ],
    },
    {
        "title": "Experience utilisateur et approche humaine",
        "paragraphs": [
            "L'approche humaine du projet consiste a concevoir l'application a partir des personnes qui l'utilisent. Un administrateur cherche a gagner du temps et a eviter les erreurs. Un professeur cherche a saisir ses notes sans parcourir des menus inutiles. Un responsable pedagogique cherche une vision claire de l'avancement.",
            "L'application reduit la fatigue administrative par la recherche, les filtres, l'import, l'export, les rapports d'erreurs et les statistiques. Elle ne remplace pas le jugement humain, mais elle elimine une partie des taches repetitives.",
            "La confiance est un autre aspect humain. Les utilisateurs doivent comprendre ce qui est sauvegarde, ce qui est en attente, ce qui est invalide et ce qui doit etre corrige. Les toasts, badges, decisions et rapports d'import participent a cette confiance.",
            "L'application respecte aussi le contexte linguistique. Le support francais/arabe et RTL n'est pas uniquement esthetique : il permet a des utilisateurs differents de travailler dans la langue la plus confortable pour eux.",
        ],
    },
    {
        "title": "Bilinguisme, RTL et accessibilite",
        "paragraphs": [
            "Le projet integre un systeme de traduction base sur un contexte de langue. La langue est stockee localement, le document HTML recoit les attributs `lang` et `dir`, et les composants adaptent leur alignement selon le mode francais ou arabe.",
            "Le mode arabe necessite une attention particuliere au RTL. Les menus, les textes, certains boutons et les axes graphiques doivent etre inverses ou alignes differemment. Cette adaptation evite une interface traduite mais visuellement incoherente.",
            "L'accessibilite peut etre encore amelioree par l'ajout de labels ARIA sur les boutons iconiques, une meilleure verification de la navigation clavier dans les modales et le controle des contrastes en mode sombre.",
            "Le bilinguisme doit aussi etre teste dans les exports et PDF. Les polices, l'encodage UTF-8 et le rendu des noms arabes doivent etre verifies avec des donnees reelles avant une utilisation officielle.",
        ],
    },
    {
        "title": "Chapitre 6 : Tests et validation",
        "paragraphs": [
            "La validation d'une application de gestion pedagogique doit couvrir les aspects fonctionnels, techniques et humains. Les tests ne se limitent pas a verifier que les pages s'affichent ; ils doivent confirmer que les regles metier sont respectees.",
            "Les tests existants couvrent principalement des aspects de base comme l'authentification. Pour renforcer le projet, il est recommande d'ajouter des tests specifiques : acces professeur, saisie de notes, import de fichiers, refus des notes invalides, gestion des rattrapages et separation des roles.",
            "Les tests manuels restent importants pour l'UI/UX. Il faut verifier la lisibilite des tableaux, la responsivite mobile, le mode sombre, le mode arabe, les modales, les imports et les messages d'erreur.",
            "Une validation avec des utilisateurs reels serait une etape tres utile. Elle permettrait de savoir si les libelles sont clairs, si les workflows correspondent aux habitudes et si certaines actions doivent etre simplifiees.",
        ],
        "bullets": [
            "Tester les droits d'acces par role.",
            "Tester les imports avec fichiers corrects et incorrects.",
            "Tester la saisie des notes entre 0 et 20 et la valeur 99.",
            "Tester le mode arabe et RTL.",
            "Tester les exports Excel et PDF.",
        ],
    },
    {
        "title": "Difficultes rencontrees et solutions apportees",
        "paragraphs": [
            "La premiere difficulte concerne la complexite metier. Les notions d'inscription pedagogique, inscription aux examens, session normale, rattrapage et decision finale doivent etre bien separees. Une confusion entre ces concepts peut conduire a des erreurs dans l'application.",
            "La deuxieme difficulte concerne les imports Excel. Les fichiers peuvent avoir des colonnes manquantes, des encodages differents, des doublons ou des valeurs invalides. La solution a consiste a ajouter des controles, des nettoyages de donnees et des rapports d'erreurs.",
            "La troisieme difficulte concerne l'interface professeur. Il fallait eviter une page trop chargee tout en affichant suffisamment d'informations. La solution a ete de proposer un dashboard minimal, une priorite actuelle et une progression paginee des modules.",
            "La quatrieme difficulte concerne le mode arabe. Le RTL, les noms arabes et les exports PDF demandent des validations specifiques. L'application integre deja des mecanismes d'adaptation, mais cette partie doit rester un point de verification continue.",
        ],
    },
    {
        "title": "Limites actuelles du projet",
        "paragraphs": [
            "Comme tout projet de stage, ProfSpace presente des limites. Certaines fonctionnalites sont deja avancees, mais des ameliorations restent possibles avant une mise en production complete.",
            "La premiere limite concerne la couverture de tests. Les tests metier doivent etre renforces afin de securiser les cas critiques : notes, rattrapages, imports et roles. La deuxieme concerne la traçabilite : un journal d'audit des modifications de notes serait tres utile.",
            "La troisieme limite concerne l'encodage et le rendu arabe. Certains textes doivent etre verifies dans tous les contextes : interface, export Excel et PDF. La quatrieme concerne la documentation utilisateur, qui doit accompagner le deploiement.",
            "Ces limites ne remettent pas en cause la valeur du projet. Elles montrent plutot les prochaines etapes logiques pour transformer une application fonctionnelle en solution institutionnelle robuste.",
        ],
    },
    {
        "title": "Chapitre 7 : Apports du stage",
        "paragraphs": [
            "Ce stage a permis de consolider des competences techniques importantes. Le developpement avec Laravel, React, Inertia.js et Tailwind CSS a permis de travailler sur une architecture moderne, proche des pratiques professionnelles.",
            "Le projet a aussi renforce les competences en modelisation. Il fallait representer correctement les donnees pedagogiques et leurs relations. Cette etape a montre l'importance de comprendre le domaine avant d'ecrire le code.",
            "Sur le plan UI/UX, le stage a permis de comprendre qu'une interface doit etre construite pour des personnes reelles. Les couleurs, la disposition, les messages et les raccourcis ne sont pas des details ; ils influencent directement la facilite d'utilisation.",
            "Enfin, le stage a developpe des competences humaines : ecoute, adaptation, communication, patience et responsabilite. Travailler sur une application liee aux notes et aux donnees d'etudiants oblige a etre rigoureux.",
        ],
    },
    {
        "title": "Perspectives d'evolution",
        "paragraphs": [
            "Plusieurs perspectives peuvent etre envisagees. A court terme, il serait utile de renforcer les tests, corriger les eventuels problemes d'encodage, finaliser la documentation utilisateur et securiser davantage la sauvegarde des notes.",
            "A moyen terme, l'application pourrait integrer un journal d'audit, des notifications, des tableaux de bord par filiere, des recherches avancees et des permissions plus fines. Ces evolutions renforceraient le pilotage administratif.",
            "A long terme, ProfSpace pourrait evoluer vers une plateforme plus large incluant un portail etudiant, des workflows de validation, des signatures numeriques et une integration avec d'autres systemes institutionnels.",
            "Ces perspectives montrent que le projet peut servir de base solide. Il ne s'agit pas d'un prototype isole, mais d'une fondation pouvant etre enrichie progressivement selon les besoins de l'etablissement.",
        ],
    },
    {
        "title": "Conclusion generale",
        "paragraphs": [
            "Le projet ProfSpace a permis de concevoir et developper une application web de gestion pedagogique adaptee au contexte de la Faculte Charia de Fes. Il couvre plusieurs besoins essentiels : gestion des donnees, inscriptions, examens, notes, statistiques et espace professeur.",
            "L'application repose sur une architecture moderne combinant Laravel, Inertia.js, React et Tailwind CSS. Elle integre des fonctionnalites avancees comme l'import/export Excel, la generation PDF, les tableaux de bord, le bilinguisme et le mode RTL.",
            "Au-dela de la technique, ce projet montre l'importance de l'approche humaine. Une bonne application administrative doit simplifier le quotidien, reduire les erreurs, rassurer les utilisateurs et respecter leur langue de travail.",
            "Ce stage a constitue une experience formatrice, a la fois technique et professionnelle. Il a permis de transformer des competences acquises en formation en une solution concrete, utile et evolutive.",
        ],
    },
    {
        "title": "Bibliographie et webographie",
        "paragraphs": [
            "Les sources suivantes ont ete utilisees pour comprendre le contexte institutionnel et les technologies mobilisees. Les informations institutionnelles exactes, comme les noms des responsables, les logos officiels et l'organigramme, doivent etre validees avec l'etablissement d'accueil avant la version finale.",
        ],
        "bullets": [
            "Laravel Documentation : https://laravel.com/docs/12.x",
            "Inertia.js Documentation : https://inertiajs.com/",
            "React Documentation : https://react.dev/learn",
            "Tailwind CSS Documentation : https://tailwindcss.com/docs/installation/framework-guides/laravel/vite",
            "Universite Mohammed V : https://www.um5.ac.ma/",
            "EST Sale : https://ests.um5.ac.ma/",
            "Universite Sidi Mohamed Ben Abdellah : https://www.usmba.ac.ma/",
            "Faculte Charia de Fes : http://sharia.usmba.ac.ma/",
        ],
    },
    {
        "title": "Annexes",
        "paragraphs": [
            "Les annexes permettent d'ajouter les elements visuels et techniques qui completent le rapport sans alourdir les chapitres principaux. Elles doivent etre completees avec les captures reelles de l'application, les diagrammes et les exemples de fichiers.",
            "Il est recommande d'ajouter des captures pour le tableau de bord administrateur, le tableau de bord professeur, la page de saisie des notes, la gestion des etudiants, la gestion des modules, l'import Excel et le releve PDF.",
            "Les annexes peuvent aussi contenir un extrait du modele de donnees, un exemple de fichier d'import, un rapport d'erreurs et un guide utilisateur court pour le professeur.",
            "Les espaces ci-dessous sont reserves aux captures et schemas finaux.",
        ],
        "placeholder": "[Captures d'ecran, schemas, exemples d'exports et signatures a inserer]",
    },
]


def add_report_pages(document):
    for index, page in enumerate(REPORT_PAGES, start=1):
        page_title(document, page["title"])
        for text in page.get("paragraphs", []):
            paragraph(document, text)
        for text in page.get("bullets", []):
            bullet(document, text)
        if page.get("placeholder"):
            placeholder_box(document, page["placeholder"], 4)
        # The preliminary pages already add six page breaks. Limiting explicit
        # breaks here keeps the final report close to the requested 30 pages
        # while allowing the last sections to flow naturally.
        if index <= 23:
            document.add_page_break()


def build():
    document = Document()
    configure(document)
    cover(document)
    add_preliminary_pages(document)
    add_report_pages(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
