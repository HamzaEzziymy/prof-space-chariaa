from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("Rapport_detaille_etapes_projet_ProfSpace.md")
OUTPUT = Path("Rapport_detaille_etapes_projet_ProfSpace.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = RGBColor(*color)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 14, RGBColor(47, 117, 181)),
        ("Heading 3", 12, RGBColor(68, 68, 68)),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_cover(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Rapport detaille des etapes de travail du projet ProfSpace")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Analyse du projet, qualite UI/UX et approche humaine")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(90, 90, 90)

    document.add_paragraph()

    table = document.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Formation", "3eme annee - Ingenierie des Applications Web et Mobiles - EST Sale"),
        ("Etablissement d'accueil", "Faculte Charia de Fes"),
        ("Stagiaire", "[Nom et prenom du stagiaire]"),
        ("Encadrant pedagogique", "[Nom de l'encadrant pedagogique]"),
        ("Encadrant professionnel", "[Nom de l'encadrant professionnel]"),
        ("Projet", "ProfSpace - Application de gestion pedagogique et de suivi des examens"),
        ("Periode du stage", "[Date de debut] - [Date de fin]"),
    ]

    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], "D9EAF7")
        set_cell_text(row.cells[0], label, bold=True, color=(31, 78, 121))
        set_cell_text(row.cells[1], value)

    document.add_page_break()


def add_markdown_line(document, line):
    stripped = line.strip()
    if not stripped or stripped == "---":
        return

    if stripped.startswith("# "):
        text = stripped[2:].strip()
        if "Rapport detaille des etapes" in text:
            return
        document.add_heading(text, level=1)
        return

    if stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=1)
        return

    if stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=2)
        return

    if stripped.startswith("#### "):
        document.add_heading(stripped[5:].strip(), level=3)
        return

    if stripped.startswith("- "):
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(stripped[2:].strip())
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        return

    if len(stripped) > 3 and stripped[0].isdigit() and ". " in stripped[:5]:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(stripped.split(". ", 1)[1].strip())
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(stripped)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def build():
    document = Document()
    configure_document(document)
    add_cover(document)

    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        add_markdown_line(document, line)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
